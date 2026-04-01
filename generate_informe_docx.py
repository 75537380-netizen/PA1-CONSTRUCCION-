from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUTPUT = "INFORME_TECNICO_UCCI_CORREGIDO.docx"


def shade_cell(cell, fill="F2F2F2"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "808080")


def add_toc(paragraph):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "Actualiza este índice en Word con clic derecho > Actualizar campo."
    fld_char3 = OxmlElement("w:fldChar")
    fld_char3.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    run._r.append(text)
    run._r.append(fld_char3)


def add_page_field(paragraph):
    run = paragraph.add_run("Página ")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_text)
    run._r.append(fld_end)


def add_center(doc, text, bold=False, size=12, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)


def add_capture_placeholder(doc, title, hint):
    p = doc.add_paragraph()
    p.add_run(title).bold = True
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Cm(15.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    shade_cell(cell, "F8F8F8")
    set_cell_border(cell)
    for _ in range(5):
        cp = cell.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.add_run(" ")
    cp = cell.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cp.add_run("[ Insertar captura aquí ]")
    r.bold = True
    r.font.size = Pt(12)
    cp2 = cell.add_paragraph()
    cp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = cp2.add_run(hint)
    r2.italic = True
    r2.font.size = Pt(10)
    for _ in range(4):
        cp = cell.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.add_run(" ")
    doc.add_paragraph()


doc = Document()

section = doc.sections[0]
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(3)
section.right_margin = Cm(2.5)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(11)
for style_name in ["Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"]:
    styles[style_name].font.name = "Calibri"

styles["Title"].font.size = Pt(22)
styles["Title"].font.bold = True
styles["Subtitle"].font.size = Pt(14)
styles["Heading 1"].font.size = Pt(15)
styles["Heading 1"].font.bold = True
styles["Heading 1"].font.color.rgb = RGBColor(31, 78, 121)
styles["Heading 2"].font.size = Pt(12)
styles["Heading 2"].font.bold = True
styles["Heading 2"].font.color.rgb = RGBColor(55, 55, 55)
styles["Heading 3"].font.size = Pt(11)
styles["Heading 3"].font.bold = True

# Carátula
for _ in range(5):
    doc.add_paragraph()
add_center(doc, "UNIVERSIDAD CONTINENTAL", bold=True, size=18, space_after=10)
add_center(doc, "INFORME TÉCNICO", bold=True, size=22, space_after=4)
add_center(doc, "Desarrollo Ágil de Aplicación con Inteligencia Artificial Generativa", bold=True, size=15, space_after=20)
add_center(doc, "Proyecto: Gestor de tareas con IA para estudiantes universitarios", size=13, space_after=20)

info_table = doc.add_table(rows=6, cols=2)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
info_table.style = "Table Grid"
info = [
    ("Curso", "[Completar]"),
    ("Docente", "[Completar]"),
    ("Integrantes", "[Completar de 3 a 5 integrantes]"),
    ("Líder del equipo", "[Completar]"),
    ("Ciclo / Sección", "[Completar]"),
    ("Fecha", "30/03/2026"),
]
for i, (k, v) in enumerate(info):
    info_table.cell(i, 0).text = k
    info_table.cell(i, 1).text = v
    shade_cell(info_table.cell(i, 0), "D9EAF7")

for _ in range(8):
    doc.add_paragraph()
add_center(doc, "Lima - Perú", size=12, space_after=0)
doc.add_page_break()

# Índice
doc.add_paragraph("Índice", style="Heading 1")
toc_p = doc.add_paragraph()
add_toc(toc_p)
doc.add_page_break()


def h1(text):
    doc.add_paragraph(text, style="Heading 1")


def h2(text):
    doc.add_paragraph(text, style="Heading 2")


h1("1. Introducción")
doc.add_paragraph(
    "El presente informe técnico describe el desarrollo de una aplicación sencilla orientada a estudiantes universitarios, cuyo objetivo es apoyar la organización académica mediante la gestión de tareas y el uso de inteligencia artificial generativa para sugerir planes de trabajo."
)
doc.add_paragraph(
    "Para esta primera entrega, el proyecto se presenta principalmente como un prototipo funcional centrado en la experiencia de usuario y el frontend. Aunque la implementación técnica contempla componentes adicionales, el alcance documentado en este informe se enfoca en la funcionalidad observable por el usuario final."
)

h1("2. Descripción del problema")
doc.add_paragraph("Muchos estudiantes universitarios tienen dificultades para organizar sus actividades académicas y personales. Entre los principales problemas identificados se encuentran:")
for item in [
    "Falta de priorización de tareas.",
    "Olvido de fechas importantes.",
    "Dificultad para visualizar tiempos disponibles.",
    "Desorganización entre actividades académicas y personales.",
    "Baja productividad por ausencia de planificación concreta.",
]:
    add_bullet(doc, item)
doc.add_paragraph("Ante este contexto, se planteó la necesidad de una solución simple, accesible y práctica que permita al estudiante registrar sus pendientes y recibir apoyo para organizar mejor su tiempo.")

h1("3. Solución propuesta")
doc.add_paragraph("La solución propuesta consiste en una aplicación web de gestión de tareas con apoyo de inteligencia artificial, diseñada específicamente para estudiantes universitarios. La aplicación permite:")
for item in [
    "Registrar tareas académicas o personales.",
    "Asignar prioridad a cada tarea.",
    "Organizar tareas según estado o categoría.",
    "Definir un horario base del estudiante.",
    "Generar un plan de trabajo asistido por IA según los pendientes registrados.",
    "Obtener sugerencias sobre huecos disponibles para avanzar tareas.",
]:
    add_bullet(doc, item)
doc.add_paragraph("La propuesta busca que el estudiante no solo anote sus actividades, sino que también reciba una orientación práctica para ejecutarlas, mejorando su administración del tiempo y reduciendo la sobrecarga mental.")

add_capture_placeholder(doc, "Espacio para captura 1", "Pantalla principal del sistema de tareas")

h1("4. Alcance de la primera entrega")
doc.add_paragraph("Para esta primera fase del proyecto se trabajó con un enfoque de prototipo funcional orientado al frontend, tomando como base los siguientes alcances:")
for item in [
    "Interfaz funcional para visualizar tareas.",
    "Formulario para registrar y editar tareas.",
    "Gestión visual de prioridades, estados y categorías.",
    "Módulo de planificación con inteligencia artificial.",
    "Apartado de horario para que el estudiante registre sus bloques fijos.",
    "Flujo de interacción enfocado en usabilidad, claridad y apoyo al estudiante.",
]:
    add_bullet(doc, item)
doc.add_paragraph("En esta entrega, la solución se presenta conceptualmente como una aplicación enfocada en la capa visible para el usuario, es decir, la experiencia de uso y el comportamiento funcional principal.")

h1("5. Metodología ágil utilizada")
doc.add_paragraph("Se decidió utilizar Scrum como metodología ágil de trabajo, debido a que permite organizar el desarrollo en iteraciones cortas, priorizar funcionalidades y validar avances parciales con rapidez.")

h2("5.1 Justificación del uso de Scrum")
for item in [
    "Facilita dividir el trabajo en entregables pequeños y funcionales.",
    "Permite adaptar el desarrollo conforme aparecen nuevas necesidades.",
    "Favorece la colaboración del equipo.",
    "Ayuda a mantener visibilidad del avance del proyecto.",
    "Es apropiado para proyectos académicos con tiempo limitado.",
]:
    add_bullet(doc, item)

h2("5.2 Organización del equipo")
for item in [
    "Líder del equipo: responsable de coordinar tareas, revisar avances y consolidar la entrega.",
    "Desarrolladores: responsables de la implementación de la interfaz, lógica funcional y documentación.",
    "Apoyo técnico / documentación: encargado de consolidar evidencias, justificar decisiones y apoyar en pruebas.",
]:
    add_bullet(doc, item)

h2("5.3 Product Backlog inicial")
for item in [
    "Como estudiante, quiero registrar tareas para no olvidar mis pendientes.",
    "Como estudiante, quiero asignar prioridad a mis tareas para saber qué hacer primero.",
    "Como estudiante, quiero visualizar mi lista de tareas de forma clara.",
    "Como estudiante, quiero registrar mi horario semanal para identificar tiempos disponibles.",
    "Como estudiante, quiero que la IA me sugiera un plan de trabajo según mis tareas.",
    "Como estudiante, quiero recibir una orientación práctica para aprovechar mis huecos libres.",
]:
    add_number(doc, item)

h2("5.4 Iteración Scrum aplicada")
doc.add_paragraph("Sprint 1: Prototipo funcional del gestor de tareas con IA", style="Heading 3")
doc.add_paragraph("Objetivo del sprint: Construir una primera versión usable de la aplicación que permita al estudiante gestionar tareas y recibir apoyo de IA para organizarse.")
doc.add_paragraph("Duración estimada: 1 semana académica.")

sprint_table = doc.add_table(rows=1, cols=2)
sprint_table.style = "Table Grid"
sprint_table.alignment = WD_TABLE_ALIGNMENT.CENTER
sprint_table.rows[0].cells[0].text = "Actividades del sprint"
sprint_table.rows[0].cells[1].text = "Entregables"
shade_cell(sprint_table.rows[0].cells[0], "D9EAF7")
shade_cell(sprint_table.rows[0].cells[1], "D9EAF7")
activities = [
    "Definición del problema y alcance",
    "Diseño de la estructura general de la aplicación",
    "Implementación de la vista principal de tareas",
    "Implementación del formulario de registro y edición",
    "Construcción del módulo de planificación con IA",
    "Incorporación del apartado de horario del estudiante",
    "Revisión de manejo de errores y mensajes de retroalimentación",
    "Redacción del informe técnico",
]
deliverables = [
    "Aplicación funcional base",
    "Interfaz de tareas operativa",
    "Integración inicial con IA generativa",
    "Flujo de planificación visible para el usuario",
    "Informe técnico de la primera entrega",
]
row = sprint_table.add_row().cells
row[0].text = "\n".join("• " + a for a in activities)
row[1].text = "\n".join("• " + d for d in deliverables)

h2("5.5 Eventos Scrum considerados")
for item in [
    "Sprint Planning: definición del problema, alcance y tareas principales.",
    "Daily Scrum informal: seguimiento del avance del equipo y ajuste de pendientes.",
    "Sprint Review: revisión de la funcionalidad construida.",
    "Sprint Retrospective: identificación de mejoras, especialmente en usabilidad, claridad visual y flujo de interacción.",
]:
    add_bullet(doc, item)

add_capture_placeholder(doc, "Espacio para captura 2", "Tablero, flujo de tareas o evidencia de la iteración Scrum")

h1("6. Uso de inteligencia artificial generativa")
doc.add_paragraph("En este proyecto se emplearon principalmente las herramientas Claude y Codex como apoyo durante el desarrollo.")
for item in [
    "Claude: para la generación y refinamiento del módulo de planificación asistida por IA.",
    "Codex: para apoyo en programación, revisión de estructura, mejoras de interfaz, redacción técnica y soporte de desarrollo.",
]:
    add_bullet(doc, item)

h2("6.1 Justificación del uso de IA")
for item in [
    "Acelerar la construcción de componentes funcionales.",
    "Mejorar la redacción y estructura de ciertos bloques de código.",
    "Proponer mejoras de interfaz y experiencia de usuario.",
    "Apoyar la generación de prompts y lógica de planificación.",
    "Facilitar tareas de revisión, depuración y documentación técnica.",
]:
    add_bullet(doc, item)

h2("6.2 Forma en que se utilizó")
for item in [
    "Sugerencia de estructuras iniciales de código.",
    "Refinamiento de componentes visuales y formularios.",
    "Mejora del flujo del módulo de planificación.",
    "Propuesta de manejo de errores más claro para el usuario.",
    "Apoyo en documentación y justificación técnica.",
]:
    add_bullet(doc, item)

h2("6.3 Impacto en el desarrollo")
doc.add_paragraph("El impacto principal fue positivo, ya que permitió reducir tiempo en tareas repetitivas y concentrar el esfuerzo del equipo en adaptar la solución al problema planteado. También ayudó a mejorar la claridad del producto final y a incorporar funcionalidades útiles para el usuario.")

h1("7. Manejo de excepciones y robustez de la aplicación")
h2("7.1 Validación de entradas del usuario")
for item in [
    "Verificación de campos obligatorios.",
    "Validación de longitud mínima de contraseña.",
    "Confirmación de contraseña en el registro.",
    "Validación del título de las tareas.",
    "Restricción de prioridades válidas.",
    "Conversión segura de fechas.",
]:
    add_bullet(doc, item)

h2("7.2 Manejo de errores en operaciones funcionales")
for item in [
    "Registro de usuario.",
    "Inicio de sesión.",
    "Creación, edición y eliminación de tareas.",
    "Gestión de categorías.",
    "Guardado de horario.",
    "Generación del plan con IA.",
    "Continuación del chat con IA.",
    "Exportación de información.",
]:
    add_bullet(doc, item)
doc.add_paragraph("La aplicación utiliza bloques try/except en los flujos principales para evitar que un error detenga por completo la experiencia del usuario. En lugar de interrumpir la ejecución, se muestran mensajes de retroalimentación claros.")

h2("7.3 Manejo de errores relacionados con IA")
for item in [
    "Error de autenticación.",
    "Error de conexión.",
    "Límite de uso alcanzado.",
    "Error devuelto por la API.",
    "Error inesperado general.",
]:
    add_bullet(doc, item)

h2("7.4 Aporte a la robustez")
for item in [
    "La aplicación no se cae ante errores comunes.",
    "Mantiene una experiencia comprensible para el usuario.",
    "Es más estable frente a entradas inválidas o fallos externos.",
    "Cumple con el criterio de manejo de excepciones solicitado en la actividad.",
]:
    add_bullet(doc, item)

h1("8. Aplicación de principios de código limpio")
h2("8.1 Separación de responsabilidades")
for item in [
    "Archivo principal para rutas y lógica de interacción.",
    "Modelos de datos definidos de forma separada.",
    "Capa de acceso a datos separada.",
    "Plantillas HTML organizadas por vistas.",
]:
    add_bullet(doc, item)

h2("8.2 Nombres descriptivos")
doc.add_paragraph("Se emplean nombres claros para funciones, variables y rutas, por ejemplo: generar_plan, guardar_horario, _parsear_fecha, _construir_descripcion_tareas y _obtener_cliente_ia.")

h2("8.3 Modularidad y reutilización")
for item in [
    "Funciones auxiliares para validación y construcción de contexto.",
    "Plantillas reutilizables para formularios.",
    "Métodos de modelo para construir objetos desde datos.",
    "Componentes visuales compartidos entre vistas.",
]:
    add_bullet(doc, item)

h2("8.4 Encapsulamiento de validaciones")
doc.add_paragraph("Parte de la validación se centraliza en los modelos, lo que permite que la lógica de consistencia no dependa únicamente de una vista o formulario específico.")

h2("8.5 Claridad estructural")
for item in [
    "Comentarios breves y útiles.",
    "Docstrings en varias funciones.",
    "Bloques seccionados por responsabilidades.",
    "Uso de tipos y parámetros opcionales en funciones.",
]:
    add_bullet(doc, item)

h1("9. Buenas prácticas de estilo de código")
for heading, items in [
    ("9.1 Consistencia en la organización", [
        "Estructura coherente entre rutas, validaciones, funciones auxiliares, plantillas y modelos.",
    ]),
    ("9.2 Mensajes claros para el usuario", [
        "Mensajes comprensibles para operaciones exitosas, advertencias y errores.",
    ]),
    ("9.3 Tipado y legibilidad", [
        "Uso de anotaciones de tipo en varias funciones para mejorar la comprensión del código.",
    ]),
    ("9.4 Interfaz pensada para usabilidad", [
        "Diseño claro y ordenado.",
        "Separación entre tareas y horario.",
        "Formulario de creación mediante diálogo.",
        "Retroalimentación inmediata.",
        "Visualización de prioridades y estados.",
    ]),
    ("9.5 Programación defensiva", [
        "Validación previa de datos, valores por defecto y control de excepciones como base de confiabilidad.",
    ]),
]:
    h2(heading)
    for item in items:
        add_bullet(doc, item)

h1("10. Refactorización aplicada")
for heading, items in [
    ("10.1 Reorganización de componentes", [
        "Separación de responsabilidades entre lógica de aplicación, modelos, acceso a datos y vistas.",
    ]),
    ("10.2 Extracción de funciones auxiliares", [
        "Obtención de categorías.",
        "Parseo de fechas.",
        "Construcción del contexto para IA.",
        "Re-renderizado del inicio para mantener el diálogo abierto.",
    ]),
    ("10.3 Reutilización de plantillas", [
        "Extracción de componentes reutilizables del formulario de tareas para evitar duplicación de código.",
    ]),
    ("10.4 Mejora del flujo de interfaz", [
        "Creación de tareas mediante un diálogo.",
        "Separación del apartado de horario del módulo de IA para una experiencia más coherente.",
    ]),
]:
    h2(heading)
    for item in items:
        add_bullet(doc, item)

h2("10.5 Refactorización orientada a claridad")
doc.add_paragraph("El objetivo principal de la refactorización no fue solo hacer que funcione, sino lograr un código más entendible, mantenible y alineado con buenas prácticas.")

add_capture_placeholder(doc, "Espacio para captura 3", "Módulo de planificación con IA o pantalla del horario del estudiante")

h1("11. Resultado funcional del proyecto")
doc.add_paragraph("Como resultado de esta primera iteración, se obtuvo una aplicación funcional que permite al estudiante:")
for item in [
    "Registrar y organizar sus tareas.",
    "Visualizar prioridades y estados.",
    "Administrar un horario base.",
    "Generar sugerencias de planificación mediante IA.",
    "Interactuar con una interfaz clara y orientada a productividad.",
]:
    add_bullet(doc, item)
doc.add_paragraph("El producto cumple con el propósito académico de demostrar el uso de metodologías ágiles, inteligencia artificial generativa, manejo de excepciones, código limpio, buenas prácticas de estilo y refactorización.")

h1("12. Conclusiones")
for item in [
    "Se logró desarrollar una aplicación sencilla pero funcional orientada a resolver un problema real de organización en estudiantes universitarios.",
    "El uso de Scrum permitió ordenar el trabajo en una iteración concreta con entregables funcionales.",
    "La integración de IA generativa mediante Claude y Codex aportó velocidad, apoyo técnico y mejora en la construcción del producto.",
    "El manejo de excepciones implementado contribuye a la robustez y estabilidad de la aplicación.",
    "La organización del código evidencia aplicación de principios de código limpio, buenas prácticas de estilo y refactorización.",
    "La solución desarrollada constituye una base sólida para futuras mejoras.",
]:
    add_number(doc, item)

h1("13. Recomendaciones")
for item in [
    "Continuar con nuevas iteraciones Scrum para ampliar la solución.",
    "Mejorar la personalización del horario por días y bloques.",
    "Incorporar métricas de productividad del estudiante.",
    "Profundizar la interacción con IA para sugerencias más específicas.",
    "Mantener la práctica de refactorización continua durante futuras mejoras.",
]:
    add_bullet(doc, item)

h1("14. Evidencias y capturas")
doc.add_paragraph("En esta sección se pueden insertar evidencias finales del funcionamiento del sistema.")
add_capture_placeholder(doc, "Espacio para captura 4", "Vista general del sistema")
add_capture_placeholder(doc, "Espacio para captura 5", "Formulario de nueva tarea o edición")
add_capture_placeholder(doc, "Espacio para captura 6", "Plan generado por IA")

h1("15. Referencias")
for item in [
    "Schwaber, K., y Sutherland, J. The Scrum Guide.",
    "Martin, R. C. Clean Code: A Handbook of Agile Software Craftsmanship.",
    "Documentación de herramientas de inteligencia artificial generativa utilizadas en programación.",
    "Material de apoyo sobre manejo de excepciones en aplicaciones web desarrolladas en Python.",
]:
    add_bullet(doc, item)

for sec in doc.sections:
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_field(footer)

doc.save(OUTPUT)
print(OUTPUT)
